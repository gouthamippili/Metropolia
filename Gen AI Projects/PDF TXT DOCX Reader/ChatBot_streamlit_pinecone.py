# Preprocessing

from langchain.llms import HuggingFaceHub
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.document_loaders import DirectoryLoader
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document as DocxDocument  # python-docx
from langchain.schema import Document as LangchainDocument  # langchain
from langchain_community.vectorstores import FAISS
import os
from dotenv import load_dotenv
from langchain.chains import RetrievalQA
from langchain_text_splitters import TextSplitter, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader

import streamlit as st
from fastapi import UploadFile

# Loading environment variables
load_dotenv()


# Load documents
def loaddata(input_type, input_data):
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
        #documents = [LangchainDocument(page_content=loader.load())]
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = [LangchainDocument(page_content=text)]
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = DocxDocument(BytesIO(input_data.read()))
        elif isinstance(input_data, UploadFile):
            doc = DocxDocument(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = [LangchainDocument(page_content=text)]
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
            documents = [LangchainDocument(page_content=text)]
        elif isinstance(input_data, UploadFile):
            text = input_data.read().decode('utf-8')
            documents = [LangchainDocument(page_content=text)]
        else:
            raise ValueError("Invalid input data for TXT")
        documents = [LangchainDocument(page_content=text)]
    else:
        raise ValueError("Unsupported input type")
    
    # Split the file
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=20)
    
    if input_type == "Link":
        texts = splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = splitter.split_documents(documents)

    # Embeddings
    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )

    # Add documents to Pinecone
    #doc_db = pc.from_documents(documents=texts, embeddings=hf_embeddings, index_name=index_name)
    doc_db= FAISS.from_documents(texts, hf_embeddings)
    return doc_db

def runQuery(query, doc_db):
    llm = HuggingFaceHub(
        repo_id="mistralai/Mistral-7B-Instruct-v0.2",
        huggingfacehub_api_token='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'
    )
    from langchain.prompts import PromptTemplate
    
    custom_template = """
    Answer the following question based on the context provided. 
    Give a direct and concise answer.

    Context: {context}
    Question: {question}
    Answer: """
    
    prompt = PromptTemplate(
        template=custom_template,
        input_variables=['context', 'question']
    )
    
    qa = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=doc_db.as_retriever(),
        return_source_documents=False,
        verbose=False,
        chain_type="stuff",
        chain_type_kwargs={
            "prompt": prompt,
        }
    )
    answer = qa({"query": query})
    return answer

def main():
    st.title("Welcome! Chat with your Documents")
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])

    input_data = None
    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=6, step=1, label="Enter the number of Links (Max count:6)")
        input_data = [st.sidebar.text_input(f"URL {i+1}") for i in range(number_input)]
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    if st.button("Proceed") and input_data:
        vectorstore = loaddata(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore

    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            answer = runQuery(query, st.session_state["vectorstore"])
            st.write(answer)

if __name__ == "__main__":
    main()